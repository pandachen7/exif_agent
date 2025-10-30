# -*- coding: utf-8 -*-
from docx import Document

try:
    # 讀取原始文檔
    doc = Document("doc/exif2accessDB(EXIF轉換資料表)_2021.docx")

    # 寫入檔案而非 stdout
    with open("doc_extracted.txt", "w", encoding="utf-8") as f:
        # 讀取所有段落
        for para in doc.paragraphs:
            if para.text.strip():
                f.write(para.text + "\n")

        # 讀取所有表格
        for table in doc.tables:
            f.write("\n=== Table ===\n")
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                f.write(" | ".join(row_data) + "\n")

    print("Document extracted successfully to doc_extracted.txt")
except Exception as e:
    print(f"Error: {e}")
